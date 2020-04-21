from docx import Document
import os
import csv
from library.FileProcessBasic import FileProcessBasic
import util


# 桩号区间list
def get_INTE(table):
    INTE = []
    for i in range(len(table.rows)):
        if table.cell(i, 1).text.strip().replace("\n", "") == "里程段（m）":
            for j in range(len(table.rows) - 1):
                per_INTE = table.cell(i + j + 1, 1).text.strip().replace("\n", "")
                first_CHAI = per_INTE.split("～")[0]
                first_CHAI_prefix = first_CHAI.split("+")[0]
                first_others = first_CHAI.split("+")[1]
                second_CHAI = per_INTE.split("～")[1]
                second_CHAI_prefix = second_CHAI.split("+")[0]
                second_others = second_CHAI.split("+")[1]
                if per_INTE == "":
                    per_INTE = "无"
                if second_CHAI_prefix == "":
                    first = (first_CHAI_prefix, first_others)
                    second = (first_CHAI_prefix, second_others)
                    first_CHAI = "+".join(first)
                    second_CHAI = "+".join(second)
                    per_INTE = (first_CHAI, second_CHAI)
                    per_INTE = "～".join(per_INTE)

                INTE.append(per_INTE)
    return INTE


class Record:
    def __init__(self, docx, INTE):
        table_result = docx.tables[1]
        table_analysis = docx.tables[2]

        name = self.get_cover(docx)
        GPRF_INTE = util.parse_GPRF_INTE(name, INTE)

        GPRF_LITH, GPRF_MD, GPRF_ZPS, GPRF_HPS, GPRF_PSB, GPRF_EM = self.get_GPRF_table_result(table_result, INTE)

        GPRF_FORE, GPRF_WEA, GPRF_INTE2, GPRF_STAB, GPRF_FAUL, GPRF_PSRL = self.get_GPRF_FORE(table_analysis, INTE)

        # 未实现
        GPRF_WATE = self.get_GPRF_WATE()
        GPRF_STRE = self.get_GPRF_STRE()

        self.dict = {
            "桩号区间": GPRF_INTE,
            "岩性": GPRF_LITH,
            "密度ρ（g/cm3）": GPRF_MD,
            "纵波速度Vp（m/s）": GPRF_ZPS,
            "横波速度Vs（m/s）": GPRF_HPS,
            "泊松比σ": GPRF_PSB,
            "动态杨氏模量（GPa）": GPRF_EM,
            "预报结果描述": GPRF_FORE,
            "风化程度": GPRF_WEA,
            "地下水": GPRF_WATE,
            "完整性": GPRF_INTE2,
            "稳定性": GPRF_STAB,
            "特殊地质情况": GPRF_FAUL,
            "推测围岩级别": GPRF_PSRL,
            "设计围岩级别": GPRF_STRE,

        }

    def get_cover(self, docx):
        name = None
        for paragraph in docx.paragraphs:
            if paragraph.text.startswith("隧道名称："):
                name = paragraph.text.split("：")[1].strip()
            if paragraph.text.startswith("项目名称："):
                name = paragraph.text.split("：")[1].strip()
            if name is not None:
                return name

    # 预报结果/预报推断结果 单元格
    def get_GPRF_FORE(self, table, val):
        GPRF_FORE = ""  # 预报结果/预报推断结果
        GPRF_WEA = ""  # 风化程度
        GPRF_INTE2 = ""  # 完整性
        GPRF_STAB = ""  # 稳定性
        GPRF_FAUL = ""  # 特殊地质情况
        GPRF_PSRL = ""  # 推测围岩级别/推断围岩级别
        for i in range(len(table.rows)):
            text = table.cell(i, 0).text
            text = text.strip().split('\n')[:-1]
            text = "".join(text)
            if text == val:
                # 预报结果/预报推断结果
                for j in range(5):
                    if table.cell(3, j).text.strip() == "预报结果" or "预报推断结果":
                        GPRF_FORE = table.cell(i, j).text.strip().replace("\n", "")
                        if GPRF_FORE == "":
                            GPRF_FORE = "无"

                        # 风化程度
                        index = GPRF_FORE.find("风化")
                        wea_start1 = GPRF_FORE.rfind("，", 0, index) + 1
                        wea_start2 = GPRF_FORE.rfind("。", 0, index) + 1
                        if wea_start1 < wea_start2:
                            wea_start = wea_start2
                        else:
                            wea_start = wea_start2
                        wea_end = GPRF_FORE.find("化", wea_start) + 1
                        GPRF_WEA = GPRF_FORE[wea_start: wea_end].replace("岩体", "")
                        if GPRF_WEA == "":
                            GPRF_WEA = "无"

                        # 完整性
                        inte2_start = GPRF_FORE.find("完整性")
                        inte2_end = GPRF_FORE.find("。", inte2_start)
                        GPRF_INTE2 = GPRF_FORE[inte2_start: inte2_end].replace("和稳定性", "")
                        if GPRF_INTE2 == "":
                            GPRF_INTE2 = "无"

                        # 稳定性
                        stab_start = GPRF_FORE.find("稳定性")
                        stab_end = GPRF_FORE.find("。", stab_start)
                        GPRF_STAB = GPRF_FORE[stab_start: stab_end]
                        if GPRF_STAB == "":
                            GPRF_STAB = "无"

                        # 特殊地质情况
                        faul_start = GPRF_FORE.find("裂隙")
                        faul_end = GPRF_FORE.find("育", faul_start) + 1
                        GPRF_FAUL = GPRF_FORE[faul_start: faul_end]
                        if GPRF_FAUL == "":
                            GPRF_FAUL = "无"

                        # 推测围岩级别/推断围岩级别
                        GPRF_PSRL = table.cell(i, j + 1).text.strip().replace("\n", "")
                        if GPRF_PSRL == "":
                            GPRF_PSRL = "无"

        return GPRF_FORE, GPRF_WEA, GPRF_INTE2, GPRF_STAB, GPRF_FAUL, GPRF_PSRL

    # 表6.1探测结果 内容
    def get_GPRF_table_result(self, table, val):
        GPRF_LITH = ""  # 岩性
        GPRF_MD = ""  # 密度
        GPRF_ZPS = ""  # 纵波速度
        GPRF_HPS = ""  # 横波速度
        GPRF_PSB = ""  # 泊松比
        GPRF_EM = ""  # 动态杨氏模量
        for i in range(len(table.rows) - 1):
            text = table.cell(i + 1, 1).text.strip().replace("\n", "")
            first_CHAI = text.split("～")[0]
            first_CHAI_prefix = first_CHAI.split("+")[0]
            first_others = first_CHAI.split("+")[1]
            second_CHAI = text.split("～")[1]
            second_CHAI_prefix = second_CHAI.split("+")[0]
            second_others = second_CHAI.split("+")[1]
            if second_CHAI_prefix == "":
                first = (first_CHAI_prefix, first_others)
                second = (first_CHAI_prefix, second_others)
                first_CHAI = "+".join(first)
                second_CHAI = "+".join(second)
                text = (first_CHAI, second_CHAI)
                text = "～".join(text)
            if text == val:
                GPRF_LITH = table.cell(i + 1, 0).text
                GPRF_MD = table.cell(i + 1, 2).text
                GPRF_ZPS = table.cell(i + 1, 3).text
                GPRF_HPS = table.cell(i + 1, 4).text
                GPRF_PSB = table.cell(i + 1, 5).text
                GPRF_EM = table.cell(i + 1, 6).text

                if GPRF_LITH == "":
                    GPRF_LITH = "无"
                if GPRF_MD == "":
                    GPRF_MD = "无"
                if GPRF_ZPS == "":
                    GPRF_ZPS = "无"
                if GPRF_HPS == "":
                    GPRF_HPS = "无"
                if GPRF_PSB == "":
                    GPRF_PSB = "无"
                if GPRF_EM == "":
                    GPRF_EM = "无"
        return GPRF_LITH, GPRF_MD, GPRF_ZPS, GPRF_HPS, GPRF_PSB, GPRF_EM

    # 地下水
    def get_GPRF_WATE(self):
        return "无"

    # 设计围岩级别
    def get_GPRF_STRE(self):
        return "无"


class Processor(FileProcessBasic):
    def save(self, output, record):
        output_path = os.path.join(output, "TSP_S1S2.csv")
        header = record.dict.keys()
        util.check_output_file(output_path, header)

        with open(output_path, "a+", encoding="utf_8_sig", newline="") as f:
            w = csv.DictWriter(f, record.dict.keys())
            w.writerow(record.dict)

    def run(self, input_path, output_path):
        files_to_process = set()
        files_to_transform = set()
        for file in os.listdir(input_path):
            absolute_file_path = os.path.join(input_path, file)
            if file.endswith(".doc"):
                files_to_transform.add(absolute_file_path)
            elif file.endswith(".docx"):
                files_to_process.add(absolute_file_path)
        files_to_delete = util.batch_doc_to_docx(files_to_transform)
        files_to_process = files_to_process.union(files_to_delete)

        for file in files_to_process:
            docx = Document(file)
            INTE = get_INTE(docx.tables[1])
            for i in range(len(INTE)):
                record = Record(docx, INTE[i])
                self.save(output_path, record)
            print("提取完成" + file)

        for file in files_to_delete:
            if os.path.exists(file):
                os.remove(file)


if __name__ == "__main__":
    test = Processor()
    inputpath = "E:/Education/409iS3/task/task4"
    outputpath = "E:/Education/409iS3/task/task4"
    test.run(inputpath, outputpath)
